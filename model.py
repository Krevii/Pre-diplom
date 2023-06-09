from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
import re

document = Document("doc.docx")

document_structure = [
            "Титульник",
            "Основная часть",
            "Заключение",
            "List of used literature",
        ]

        # Примеры текстов контента для каждого раздела
section_contents = [
            ["Текст титульника"],
            ["Текст работы (проекта) следует печатать, соблюдая следующие размеры полей: \"левое – 30 мм, верхнее – 20 мм, правое – 10 мм и нижнее – 25+\" мм.Не разрешается использовать компьютерные возможности акцентирования внимания на определенных терминах, формулах, теоремах, применяя шрифты разной гарнитуры, а только курсивное начертание. Вне зависимости от способа выполнения работы (проекта) качество напечатанного текста и оформление иллюстраций, таблиц, распечаток с ПЭВМ должно удовлетворять требованию их четкого воспроизведения.Фамилии, названия учреждений, организаций, фирм, название изделий и другие имена собственные приводят на языке оригинала.Наименования структурных элементов дипломной работы (проекта) «СОДЕРЖАНИЕ», «ПЕРЕЧЕНЬ ИСПОЛЬЗУЕМЫХ ОБОЗНАЧЕНИЙ И СОКРАЩЕНИЙ», «ВВЕДЕНИЕ», «ЗАКЛЮЧЕНИЕ», «СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ» служат заголовками структурных элементов работы.Дипломную (курсовую) работу (проект) следует делить на разделы и подразделы. Каждый раздел и подраздел должен содержать законченную информацию.Нименования разделов в совокупности должны раскрывать тему работы (проекта), а наименования подразделов в совокупности должны раскрывать соответствующий раздел.Наименования разделов и подразделов должны четко и кратко отражать их содержание.Наименования разделов и подразделов следует печатать с абзацного отступа с прописной буквы без точки в конце, ни как не выделяя.Если наименование состоит из двух предложений, их разделяют точкой. Страницы работы (проекта) следует нумеровать арабскими цифрами, соблюдая сквозную нумерацию по всему тексту, с применением шрифта – Times New Roman, (кегль – 14 пунктов). Номер страницы проставляют в центре нижней части листа без точки. Обложка, титульный лист включаются в общую нумерацию страниц. Номер страницы проставляется с введения.Иллюстрации и таблицы, расположенные на отдельных листах, включают в общую нумерацию страниц работы (проекта).Разделы дипломной (курсовой) работы (проекта) должны иметь порядковые номера в пределах всего документа, обозначенные арабскими цифрами без точки и записанные с абзацного отступа. Подразделы должны иметь нумерацию в пределах каждого раздела. Номер подраздела состоит из номеров раздела и подразделзделенных точкой.", "В конце Номер подраздела состоит из номеров раздела и подразделзделенных точкой. В концеНомер подраздела состоит из номеров раздела и подразделзделенных точкой. В концеНомер подраздела состоит из номеров раздела и подразделзделенных точкой.", "В концеНомер подраздела состоит из номеров раздела и подразделзделенных точкой. В концеНомер подраздела состоит из номеров раздела и подразделзделенных точкой. В концеНомер подраздела состоит из номеров раздела и подразделзделенных точкой. В концеНомер подраздела состоит из номеров раздела и подразделзделенных точкой. В концеНомер подраздела состоит из номеров раздела и подразделзделенных точкой. В концеНомер подраздела состоит из номеров раздела и подразделзделенных точкой. В конце"],
            ["1. Количество слайдов презентации для защиты дипломной работы – 12-15. Меньшее количество не позволяет раскрыть смысл излагаемого материала, большее количество превращается в формальное перелистывание страниц. Для курсовой работы объем составляет 7-8 слайдов.2. Первый и последний слайды должны быть одинаковыми. На них указывается полное наименование учебного заведения; тема дипломной работы; фамилия, имя отчество студента; наименование получаемой специальности или факультета; фамилия, имя, отчество научного руководителя; город; год защиты.Первый слайд представляет Государственной комиссии студента его дипломную работу, последний (дублирующий первый) – позволяет членам комиссии (некоторые из которых видят соискателя диплома впервые) обратиться к нему по имени-отчеству. Не надо писать на последнем слайде: «Спасибо за внимание!». Это не воспринимается преподавателями как уважение к ним, а, скорее – как попытка уменьшить дистанцию между студентом и членами комиссии. При создании презентации к курсовой следуйте тому же правилу.3. Слайды, расположенные после первого, могут быть распределены следующим образом. На нескольких указывается актуальность, объект, предмет, проблема, цель, задачи исследования в соответствии с текстом защиты. Иногда на слайдах представляется структура дипломной (курсовой) работы, содержание глав, но это не несет смысловой нагрузки. Зато на слайдах могут указываться основные понятия, на которые опирается исследователь с обязательным указанием, откуда взята цитата и кто ее автор. Затем на слайдах представляются графики, таблицы, иллюстрирующие данные проведенного исследования или ход эксперимента. Следует обратить внимание на то, что каждый такой слайд должен иметь заголовок.4. Размер шрифта на слайдах должен быть не менее 28, иначе текст никто не увидит. Заголовки выделяются и пишутся размером шрифта не менее 36. Фон слайда желательно подобрать однотонный, не ядовитый. Цвет шрифта – темный на светлом фоне, без тени.5. Теперь о смене слайдов во время защиты. Конечно, неудобно управлять презентаций самому докл дчику, можно поручить работу с мультимедийной презентацией однокурснику. Чтобы слайды соответствовали тексту, необходимо напечатать для иллюстратора еще один экземпляр речи и разместить на нем инструкцию по смене слайдов. И, конечно, стоит 1-2 раза прорепетировать защиту вдвоем.6. Некрасиво смотрятся на большом экране орфографические и пунктуационные ошибки. Это снижает впечатление от выступления. Можно и даже нужно попросить человека, грамотности которого Вы доверяете, проверить текст.7. Возможности техники различные. Поэтому презентация, где слайды «вылетают», текст «выезжает» или появляется из ничего, может просто «зависнуть». Это один из случаев, когда показ Вашего умения работать в PowerPoint, может сыграть злую шутку на защите. Лучше подготовить простую презентацию с простой сменой слайдов.8. Презентация дипломной работы должна быть установлена на компьютер, подключенный к проектору заранее – до начала процедуры защиты дипломов всей группы. На рабочем столе создайте папку со своей фамилией,чтобы презентацию легко было найти. И проверьте до защиты, откроется ли презентация на этом оборудовании, совместима ли программа, в которой создана Ваша презентация с возможностями рабочего компьютера, стоящего в аудитории.Качественная презентация дипломной или курсовой работы, сопровождающая грамотно написанную речь, существенно увеличивает шансы получить высокую отметку на защите."],
            ["В заключение можно отметить, что компьютерные игры уже давно стали неотъемлемой частью нашей жизни, а их разработка - трудоемкой задачей, которая требует усилий и знаний специалистов из различных областей. Однако, разработка игры должна быть направлена не только на развлечение, но и на достижение определенных целей, соответствующих потребностям и интересам пользователей. В данной работе была реализована игра, которая помогает развивать навыки многозадачности, тактики и критической оценки собственных возможностей в малознакомой ситуации, при этом не обязывая погружаться в игру слишком сильно. Кроме того, стоит отметить, что разработка игр может быть выполнена для различных системных платформ, что позволяет максимально расширить аудиторию пользователей. В целом, игры могут стать не только средством развлечения, но и инструментом для обучения, общения и достижения личных целей.Одним важным аспектом в разработке компьютерных игр является использование новых технологий и инновационных подходов. С появлением виртуальной реальности, искусственного интеллекта и блокчейн-технологий, появляются новые возможности для создания уникальных игровых продуктов. Однако, необходимо помнить, что главной целью любой игры должно быть удовлетворение потребностей пользователя, а не использование новых технологий в целях самоутверждения разработчиков.Таким образом, компьютерные игры являются неотъемлемой частью современной культуры и имеют широкое применение в различных областях, от развлечений и общения до обучения и бизнеса. Важно помнить, что разработка игр является трудоемким процессом, требующим совместной работы множества специалистов и использования поэтапного подхода. В результате получается продукт, который должен удовлетворять потребности и ожидания пользователей, а также соответствовать современным технологическим требованиям."],
        ]

class Model:
    def __init__(self):
        self.data = ["somedata1", "somedata2"]

    def get_data(self):
        return self.data

    def delete_paragraph(self, paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None

    def replace_quotes(self, text):
        def replace_quote(match):
            replace_quote.open_quote = not replace_quote.open_quote
            return '»' if replace_quote.open_quote else '«'

        replace_quote.open_quote = True
        return re.sub(r'"', replace_quote, text)

    def is_image_at_page_start(self, paragraph):
        # Здесь можно добавить вашу логику для определения, находится ли изображение в начале страницы.
        # В качестве примера, мы считаем, что изображение находится в начале страницы, если перед ним нет текста.
        return len(paragraph.text.strip()) == 0

    def report_validation(self):

        for paragraph in document.paragraphs:
            self.delete_paragraph(paragraph)

        font_name = "Times New Roman"
        font_size = Pt(14)
        picture = []
        description_picture = []

        document.add_paragraph()
        # Определение структуры документа
        section_itr = 0
        content_itr = 1
        for i, paragraph in enumerate(document.paragraphs):
            paragraph.text = paragraph.text.upper()

            for j, header, content in zip(enumerate(document_structure), document_structure, section_contents ):
                content_paragraph = content
                section_header = document.add_paragraph(header)
                document.add_paragraph()
                print(content)
                for k, content_item in enumerate(content):

                    
                    picture_place = document.add_paragraph()
                    picture_place.add_run().add_picture("resources/dont care about us.jpg", width=Inches(4))
                    picture.append(picture_place)

                    document.add_paragraph()
                    description_picture.append(document.add_paragraph(
                        f"Рисунок {list(range(len(picture)))[-1]}.{content_itr}. Чоткая картинка."))
                    document.add_paragraph()
                    content_itr = content_itr + 1

                    # Добавление текста контента
                    content_paragraph = document.add_paragraph(self.replace_quotes(content[k]))

                if self.is_image_at_page_start(picture_place):
                    paragraph.add_run("Текст перед изображением на новой странице")

                # Добавление разрыва страницы перед заголовком раздела (кроме первого)
                if content_paragraph != document_structure[0] and header != document_structure[-1]:
                    content_paragraph.add_run().add_break(WD_BREAK.PAGE)

        for i, paragraph in enumerate(document.paragraphs):

            paragraph.style.font.size = font_size
            paragraph.style.font.name = font_name

            paragraph_format = paragraph.paragraph_format

            paragraph_format.left_indent = 0
            paragraph_format.right_indent = 0
            paragraph_format.space_after = 0
            paragraph_format.space_before = 0
            paragraph_format.line_spacing = 1
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph_format.first_line_indent = Pt(42.5)

        for picture_item, description_picture_item in zip(picture, description_picture):
            picture_item.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            description_picture_item.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            picture_item.paragraph_format.first_line_indent = Pt(0)
            description_picture_item.paragraph_format.first_line_indent = Pt(0)

            # paragraph.text.title()

        document.save("doc.docx")
